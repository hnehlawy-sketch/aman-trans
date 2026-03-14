#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_text.py — Extract plain text from PDF or DOCX files
Auto-installs required packages if missing.
Usage: python extract_text.py <input_file>
"""

import sys
import os
import json
import subprocess

# ── Auto-installer ─────────────────────────────────────────────────────────────
def pip_install(*packages):
    subprocess.run(
        [sys.executable, '-m', 'pip', 'install', '--quiet', '--disable-pip-version-check', *packages],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )

def ensure_packages():
    missing = []
    try:
        import docx  # noqa
    except ImportError:
        missing.append('python-docx')
    try:
        import pdfplumber  # noqa
    except ImportError:
        missing.append('pdfplumber')
    try:
        from pypdf import PdfReader  # noqa
    except ImportError:
        missing.append('pypdf')
    try:
        from reportlab.lib.pagesizes import A4  # noqa
    except ImportError:
        missing.append('reportlab')
    if missing:
        pip_install(*missing)

ensure_packages()

# ── Extraction ─────────────────────────────────────────────────────────────────

def extract_pdf(path):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[Page {i+1}]\n{text.strip()}")
        if pages:
            return "\n\n".join(pages)
    except Exception as e:
        sys.stderr.write(f"pdfplumber failed: {e}\n")

    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")
        if pages:
            return "\n\n".join(pages)
        raise RuntimeError("No text found — PDF may be image-only (scanned).")
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF: {e}")


def extract_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            style = (para.style.name or '') if para.style else ''
            if 'Heading 1' in style or style == 'Title':
                parts.append(f"# {txt}")
            elif 'Heading 2' in style:
                parts.append(f"## {txt}")
            elif 'Heading 3' in style:
                parts.append(f"### {txt}")
            else:
                parts.append(txt)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                seen = []
                for c in cells:
                    if not seen or seen[-1] != c:
                        seen.append(c)
                line = " | ".join(seen)
                if line.strip():
                    parts.append(line)

        if parts:
            return "\n\n".join(parts)
        raise RuntimeError("Document appears empty.")

    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX: {e}")


def extract_doc_legacy(path):
    try:
        result = subprocess.run(['antiword', path], capture_output=True, timeout=30)
        if result.returncode == 0:
            return result.stdout.decode('utf-8', errors='replace')
    except Exception:
        pass
    raise RuntimeError(
        "Cannot read legacy .doc file. "
        "Please convert to .docx first: open in Word > Save As > .docx"
    )


def main():
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: extract_text.py <file>"}))
        sys.exit(1)

    input_path = sys.argv[1]

    if not os.path.exists(input_path):
        print(json.dumps({"error": f"File not found: {input_path}"}))
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()

    try:
        if ext == '.pdf':
            text = extract_pdf(input_path)
            doc_type = 'pdf'
        elif ext == '.docx':
            text = extract_docx(input_path)
            doc_type = 'docx'
        elif ext == '.doc':
            text = extract_doc_legacy(input_path)
            doc_type = 'doc'
        else:
            with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            doc_type = 'text'

        word_count = len(text.split())
        print(json.dumps({
            "success": True,
            "type": doc_type,
            "wordCount": word_count,
            "text": text
        }, ensure_ascii=False))

    except Exception as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == '__main__':
    main()
